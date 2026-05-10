from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document('Doc2.docx')

image_path = 'p.png'

table = doc.add_table(rows=3, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(15):
    row = i // 5
    col = i % 5
    cell = table.rows[row].cells[col]
    cell_para = cell.paragraphs[0]
    run = cell_para.add_run()
    run.add_picture(image_path, width=Inches(1.3))

doc.save('Doc2_con_imagenes.docx')
print('Se insertaron 15 imágenes en matriz 3x5 en una hoja')